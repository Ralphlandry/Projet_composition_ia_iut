"""Generate a Word document about EvalPro platform advantages and security."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.paragraph_format.space_after = Pt(6)

style_h1 = doc.styles['Heading 1']
style_h1.font.name = 'Calibri'
style_h1.font.size = Pt(18)
style_h1.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
style_h1.font.bold = True

style_h2 = doc.styles['Heading 2']
style_h2.font.name = 'Calibri'
style_h2.font.size = Pt(14)
style_h2.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
style_h2.font.bold = True

style_h3 = doc.styles['Heading 3']
style_h3.font.name = 'Calibri'
style_h3.font.size = Pt(12)
style_h3.font.color.rgb = RGBColor(0x37, 0x4B, 0x8C)
style_h3.font.bold = True


def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(f' {text}')
    else:
        p.add_run(text)
    return p


def add_para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


# ══════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('EvalPro')
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Plateforme Intelligente de Gestion d\'Examens')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Document de Présentation — Avantages & Sécurité')
run.font.size = Pt(14)
run.bold = True

doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run(f'Date : {datetime.now().strftime("%d/%m/%Y")}')
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════
doc.add_heading('Table des Matières', level=1)
toc_items = [
    '1. Introduction',
    '2. Architecture Technique',
    '3. Avantages Fonctionnels de la Plateforme',
    '   3.1 Gestion complète des examens',
    '   3.2 Correction intelligente par IA',
    '   3.3 Système de rôles et permissions',
    '   3.4 Tableau de bord analytique avancé',
    '   3.5 Internationalisation (i18n)',
    '   3.6 Système de notifications intégré',
    '   3.7 Export de données multiformats',
    '   3.8 Gestion des classes et profils étudiants',
    '   3.9 Import PDF intelligent',
    '   3.10 Interface moderne et responsive',
    '4. Aspects de Sécurité',
    '   4.1 Authentification et gestion des sessions',
    '   4.2 Contrôle d\'accès basé sur les rôles (RBAC)',
    '   4.3 Protection contre les injections SQL',
    '   4.4 Protection contre les attaques XSS',
    '   4.5 Limitation de débit (Rate Limiting)',
    '   4.6 Isolation des données utilisateur',
    '   4.7 Protection anti-triche',
    '   4.8 Journalisation et audit',
    '   4.9 Gestion sécurisée des mots de passe',
    '   4.10 Configuration CORS',
    '   4.11 Validation des entrées',
    '   4.12 Mécanisme de désactivation de compte',
    '5. Technologies Utilisées',
    '6. Conclusion',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction', level=1)
add_para(
    'EvalPro est une plateforme intelligente de gestion d\'examens développée dans le cadre '
    'd\'un projet universitaire. Elle offre une solution complète pour la création, la gestion, '
    'la passation et la correction d\'examens en ligne, intégrant des fonctionnalités avancées '
    'd\'intelligence artificielle pour l\'auto-correction des copies.'
)
add_para(
    'Ce document présente en détail les avantages fonctionnels de la plateforme ainsi que '
    'l\'ensemble des mesures de sécurité implémentées pour garantir l\'intégrité, la '
    'confidentialité et la disponibilité des données.'
)
add_para(
    'La plateforme s\'adresse à trois profils d\'utilisateurs : les administrateurs, les '
    'professeurs et les étudiants, chacun disposant d\'un espace dédié avec des fonctionnalités '
    'adaptées à son rôle.'
)

# ══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE TECHNIQUE
# ══════════════════════════════════════════════════════════════
doc.add_heading('2. Architecture Technique', level=1)
add_para(
    'EvalPro repose sur une architecture moderne en trois couches, séparant clairement le '
    'frontend, le backend et le service d\'intelligence artificielle :'
)

# Architecture table
table = doc.add_table(rows=4, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Couche', 'Technologies', 'Port']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

data = [
    ('Frontend', 'React 18, TypeScript, Vite, TailwindCSS, shadcn/ui', '8080'),
    ('Backend API', 'FastAPI (Python), SQLAlchemy ORM, PostgreSQL', '8001'),
    ('Service IA', 'FastAPI (Python), Ollama, Qwen2.5:3b', '8000'),
]
for r, row_data in enumerate(data, 1):
    for c, val in enumerate(row_data):
        table.rows[r].cells[c].text = val

doc.add_paragraph()
add_para(
    'Cette architecture découplée permet une maintenance indépendante de chaque composant, '
    'une scalabilité horizontale et une meilleure testabilité. Le frontend communique avec le '
    'backend via une API REST sécurisée par JWT, et le backend délègue la correction '
    'intelligente au microservice IA dédié.'
)

# ══════════════════════════════════════════════════════════════
# 3. AVANTAGES FONCTIONNELS
# ══════════════════════════════════════════════════════════════
doc.add_heading('3. Avantages Fonctionnels de la Plateforme', level=1)

# 3.1
doc.add_heading('3.1 Gestion complète du cycle de vie des examens', level=2)
add_para(
    'EvalPro offre une gestion intégrale du cycle de vie d\'un examen, de sa création '
    'à la publication des résultats :'
)
add_bullet('Création d\'examens avec paramètres détaillés : titre, description, matière, spécialité, niveau, classe, semestre, type d\'évaluation, durée, barème, dates de début et fin.')
add_bullet('Système de statuts progressif : Brouillon → Programmé → Publié → Clôturé (automatique).')
add_bullet('Publication automatique des examens programmés lorsque la date de début est atteinte.')
add_bullet('Clôture automatique des copies en cours lorsque la deadline est dépassée.')
add_bullet('Structure multi-parties pour les examens complexes (parties avec titres, sous-titres, descriptions).')
add_bullet('Support de plusieurs types de questions : QCM, Vrai/Faux, réponse courte, réponse longue.')
add_bullet('Paramétrage de la difficulté et du barème par question.')

# 3.2
doc.add_heading('3.2 Correction intelligente par Intelligence Artificielle', level=2)
add_para(
    'L\'un des atouts majeurs d\'EvalPro est son système de correction automatique alimenté '
    'par un modèle de langage (LLM) exécuté localement :'
)
add_bullet('Correction déterministe pour les QCM et Vrai/Faux (correspondance exacte, par lettre, par texte d\'option).')
add_bullet('Correction IA bienveillante pour les réponses courtes avec un minimum de 20% si la réponse n\'est pas vide.')
add_bullet('Correction IA avec analyse conceptuelle pour les réponses longues (minimum 25% si effort démontré).')
add_bullet('Vérification backend systématique : les réponses QCM/VF sont re-vérifiées de façon déterministe indépendamment du résultat IA.')
add_bullet('Règle des réponses vides : toute réponse vide obtient automatiquement 0 point.')
add_bullet('Statut intermédiaire « Corrigé Auto » permettant au professeur de valider ou ajuster avant la note définitive.')
add_bullet('Suggestion automatique de réponses correctes pour les questions importées (traitement parallèle avec ThreadPoolExecutor).')
add_bullet('Modèle IA pré-chargé au démarrage du service pour des temps de réponse optimaux.')

# 3.3
doc.add_heading('3.3 Système de rôles et permissions', level=2)
add_para(
    'La plateforme implémente un contrôle d\'accès granulaire basé sur trois rôles distincts :'
)

role_table = doc.add_table(rows=4, cols=2)
role_table.style = 'Light Grid Accent 1'
role_table.alignment = WD_TABLE_ALIGNMENT.CENTER
role_table.rows[0].cells[0].text = 'Rôle'
role_table.rows[0].cells[1].text = 'Périmètre'
for p in role_table.rows[0].cells[0].paragraphs:
    p.runs[0].bold = True
for p in role_table.rows[0].cells[1].paragraphs:
    p.runs[0].bold = True

role_data = [
    ('Administrateur', 'Accès complet : gestion des utilisateurs (CRUD), matières, spécialités, niveaux, logs d\'audit, diffusion de notifications.'),
    ('Professeur', 'Création et gestion de ses propres examens, corrections, analytiques, notes, gestion de classes.'),
    ('Étudiant', 'Passage d\'examens selon sa spécialité/niveau, consultation de ses résultats, gestion de ses notifications.'),
]
for r, (role, scope) in enumerate(role_data, 1):
    role_table.rows[r].cells[0].text = role
    role_table.rows[r].cells[1].text = scope

doc.add_paragraph()
add_para(
    'Chaque route frontend est protégée par un composant ProtectedRoute qui vérifie le rôle '
    'de l\'utilisateur avant d\'autoriser l\'accès. Côté backend, chaque endpoint vérifie '
    'systématiquement le rôle via get_user_role().'
)

# 3.4
doc.add_heading('3.4 Tableau de bord analytique avancé', level=2)
add_para(
    'Le module d\'analyses offre une vision complète des performances pédagogiques :'
)
add_bullet('8 indicateurs clés (KPI) : moyenne générale, taux de réussite, nombre d\'étudiants, corrections en attente, meilleure note, plus basse note, copies notées, nombre d\'épreuves.')
add_bullet('Filtrage par matière pour des analyses ciblées.')
add_bullet('Graphique de répartition des notes en 6 tranches (0-4, 5-8, 9-11, 12-14, 15-17, 18-20).')
add_bullet('Tableau détaillé par épreuve avec moyenne, min, max et taux de réussite.')
add_bullet('Classement des 10 meilleurs étudiants avec matricule, nombre d\'épreuves et moyenne.')
add_bullet('Export Excel complet du rapport analytique.')

# 3.5
doc.add_heading('3.5 Internationalisation (i18n)', level=2)
add_para(
    'EvalPro supporte intégralement le bilinguisme français/anglais :'
)
add_bullet('Plus de 700 clés de traduction couvrant l\'intégralité de l\'interface.')
add_bullet('Système de traduction via hook useLanguage() et fonction t() appliqué à toutes les pages.')
add_bullet('Persistance de la langue par utilisateur dans le localStorage.')
add_bullet('Conservation de la préférence linguistique même après déconnexion grâce à un double stockage (clé utilisateur + clé générique).')
add_bullet('Basculement instantané sans rechargement de page.')

# 3.6
doc.add_heading('3.6 Système de notifications intégré', level=2)
add_bullet('Notifications in-app avec titre, message, type (info/succès) et statut lu/non lu.')
add_bullet('Notifications automatiques : publication d\'examen aux étudiants concernés, finalisation de note.')
add_bullet('Diffusion de notifications par l\'administrateur à des destinataires ciblés.')
add_bullet('Gestion par l\'étudiant : marquer comme lu, supprimer.')

# 3.7
doc.add_heading('3.7 Export de données multiformats', level=2)
add_bullet('Export Excel (.xlsx) des notes, rapports analytiques et listes d\'étudiants.')
add_bullet('Export PDF des relevés de notes avec mise en page professionnelle.')
add_bullet('Inclusion du matricule, de la décision (Validé/Non validé) et de toutes les métadonnées pertinentes.')
add_bullet('Formatage automatique et nommage des fichiers avec horodatage.')

# 3.8
doc.add_heading('3.8 Gestion des classes et profils étudiants', level=2)
add_bullet('Création et gestion de classes avec attribution de matières.')
add_bullet('Profils étudiants enrichis : matricule (numéro étudiant), spécialité, niveau académique.')
add_bullet('Affichage du matricule dans toutes les vues de notes, corrections et analyses.')
add_bullet('Filtrage automatique des examens disponibles selon la spécialité et le niveau de l\'étudiant.')

# 3.9
doc.add_heading('3.9 Import PDF intelligent', level=2)
add_bullet('Upload de fichiers PDF contenant des questions d\'examen.')
add_bullet('Extraction automatique des questions par analyse regex.')
add_bullet('Structuration automatique en questions typées (QCM, courte, longue).')
add_bullet('Suggestion de réponses correctes par l\'IA pour les questions importées.')

# 3.10
doc.add_heading('3.10 Interface moderne et responsive', level=2)
add_bullet('Interface React construite avec shadcn/ui et TailwindCSS pour un design professionnel et cohérent.')
add_bullet('Mode sombre / mode clair avec persistance par utilisateur.')
add_bullet('Navigation responsive avec sidebar pour desktop et barre de navigation inférieure pour mobile.')
add_bullet('Animations fluides et retours visuels (toasts, badges, indicateurs de progression).')
add_bullet('Dialogues de confirmation pour les actions critiques (suppression, soumission).')
add_bullet('Composants réutilisables : tableaux triables, filtres, sélecteurs, cartes statistiques.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. ASPECTS DE SÉCURITÉ
# ══════════════════════════════════════════════════════════════
doc.add_heading('4. Aspects de Sécurité', level=1)
add_para(
    'La sécurité est au cœur de la conception d\'EvalPro. La plateforme implémente de '
    'multiples couches de protection couvrant les principales vulnérabilités identifiées '
    'par le référentiel OWASP Top 10.'
)

# 4.1
doc.add_heading('4.1 Authentification et gestion des sessions (OWASP A07)', level=2)
add_para('Mécanismes d\'authentification :', bold=True)
add_bullet('Authentification par JWT (JSON Web Token) avec algorithme HS256 via la bibliothèque python-jose.')
add_bullet('Durée de vie configurable des tokens (par défaut 24h via JWT_EXPIRE_MINUTES).')
add_bullet('Extraction sécurisée du token depuis l\'en-tête Authorization: Bearer <token>.')
add_bullet('Messages d\'erreur distincts pour les tokens expirés (« Token expiré ») et invalides (« Token invalide ») — retour HTTP 401.')
add_bullet('Stockage du token dans sessionStorage (et non localStorage), réduisant la persistance en cas d\'attaque XSS.')
add_bullet('Déconnexion automatique en cas de réponse 401 du serveur.')
add_bullet('Blocage de l\'auto-inscription en tant qu\'administrateur.')

# 4.2
doc.add_heading('4.2 Contrôle d\'accès basé sur les rôles — RBAC (OWASP A01)', level=2)
add_para(
    'Le contrôle d\'accès est le premier élément du OWASP Top 10 (Broken Access Control). '
    'EvalPro le traite avec rigueur :'
)
add_bullet('Table user_roles associant chaque utilisateur à un rôle unique (admin, professeur, etudiant).')
add_bullet('Vérification systématique du rôle côté backend avant toute opération de lecture ou d\'écriture.')
add_bullet('Composant ProtectedRoute côté frontend n\'autorisant l\'accès qu\'aux rôles spécifiés.')
add_bullet('Les professeurs ne peuvent accéder qu\'à leurs propres examens (filtre Exam.created_by == current_user.id).')
add_bullet('Les étudiants sont strictement limités à leurs propres soumissions, réponses et notifications.')
add_bullet('Listes blanches d\'insertion et de mise à jour par rôle — les étudiants ne peuvent par exemple pas modifier leur propre note.')

# 4.3
doc.add_heading('4.3 Protection contre les injections SQL (OWASP A03)', level=2)
add_bullet('Utilisation exclusive de SQLAlchemy ORM — aucune interpolation de chaîne SQL brute dans les requêtes.')
add_bullet('Système de filtres structuré : les filtres JSON côté client sont transformés en appels ORM .filter() typés.')
add_bullet('Les seules requêtes SQL littérales (text()) sont des DDL de démarrage (CREATE TABLE IF NOT EXISTS) avec du SQL codé en dur, sans paramètre utilisateur.')
add_bullet('Validation Pydantic de toutes les entrées API avant traitement.')

# 4.4
doc.add_heading('4.4 Protection contre les attaques XSS (OWASP A03)', level=2)
add_bullet('React échappe automatiquement toutes les données rendues dans le JSX, empêchant l\'injection de scripts.')
add_bullet('Token de session stocké dans sessionStorage plutôt que localStorage (plus éphémère, limité à l\'onglet).')
add_bullet('Aucune utilisation de dangerouslySetInnerHTML dans le code de l\'application.')
add_bullet('Validation des entrées côté serveur via Pydantic avec des types stricts (EmailStr, etc.).')

# 4.5
doc.add_heading('4.5 Limitation de débit — Rate Limiting (OWASP A04)', level=2)
add_bullet('Limitation en mémoire sur la route de connexion : 5 tentatives maximum par adresse IP sur une fenêtre de 5 minutes.')
add_bullet('Protection efficace contre les attaques par force brute sur les mots de passe.')
add_bullet('Les tentatives au-delà de la limite reçoivent une réponse HTTP 429 (Too Many Requests).')

# 4.6
doc.add_heading('4.6 Isolation des données utilisateur (OWASP A01)', level=2)
add_para(
    'Chaque rôle dispose d\'un périmètre strict de données accessibles :'
)
add_bullet('Étudiants : filtrés au niveau requête pour ne voir que leurs propres soumissions, réponses, notifications et profil. Accès bloqué aux tables classes, class_students et questions.')
add_bullet('Professeurs : visibilité limitée à leurs propres examens et aux soumissions/questions associées.')
add_bullet('Les réponses correctes (correct_answer) sont systématiquement supprimées des réponses API pour les étudiants.')
add_bullet('Unicité des soumissions : un seul passage par étudiant par examen, vérifiée côté serveur.')

# 4.7
doc.add_heading('4.7 Protection anti-triche', level=2)
add_bullet('Application de la deadline côté serveur : un étudiant ne peut pas soumettre après la fin de l\'examen.')
add_bullet('Restriction de statut : un étudiant ne peut positionner le statut qu\'à « en_cours » ou « soumis », jamais à « corrigé ».')
add_bullet('Listes blanches de champs : à l\'insertion, les étudiants ne peuvent renseigner que {id, exam_id, student_id, started_at} — impossible de se donner une note.')
add_bullet('À la mise à jour, les étudiants ne peuvent modifier que {status, incidents} sur les soumissions et {answer_text} sur les réponses.')
add_bullet('Système d\'incidents : les événements anti-triche (changement d\'onglet, perte de connexion) sont enregistrés dans le champ incidents de la soumission.')
add_bullet('Notification automatique au professeur en cas de comportement suspect.')

# 4.8
doc.add_heading('4.8 Journalisation et audit (OWASP A09)', level=2)
add_para(
    'EvalPro implémente un système complet d\'audit trail pour la traçabilité :'
)
add_bullet('Table audit_logs immuable enregistrant : user_id, user_email, action (insert/update/delete), table_name, row_id, changes (JSON), ip_address, created_at.')
add_bullet('Un enregistrement d\'audit est créé automatiquement après chaque opération d\'insertion, mise à jour ou suppression dans la couche CRUD générique.')
add_bullet('Les erreurs de journalisation sont silencées pour ne pas bloquer les opérations métier, tout en préservant la robustesse du système.')
add_bullet('Interface d\'administration dédiée (AuditLogs) permettant la consultation et la recherche dans l\'historique des actions.')

# 4.9
doc.add_heading('4.9 Gestion sécurisée des mots de passe (OWASP A02)', level=2)
add_bullet('Hachage bcrypt via passlib.CryptContext avec protection automatique contre les mots de passe dépassant 72 octets.')
add_bullet('Mot de passe minimum de 6 caractères à l\'inscription, 8 caractères pour les réinitialisations administrateur.')
add_bullet('Les mots de passe ne sont jamais stockés en clair ni transmis dans les réponses API.')
add_bullet('Mécanisme de désactivation de compte par préfixe <disabled> sur le hash — empêchant toute connexion sans supprimer le compte.')
add_bullet('Impossibilité pour un administrateur de se désactiver lui-même (protection contre le verrouillage).')

# 4.10
doc.add_heading('4.10 Configuration CORS (OWASP A05)', level=2)
add_bullet('Configuration CORS via la variable d\'environnement CORS_ORIGINS permettant de lister les origines autorisées.')
add_bullet('En mode production, seules les origines spécifiquement listées sont autorisées.')
add_bullet('Les credentials sont désactivées lorsque l\'origine est définie en mode wildcard (*).')
add_bullet('En-têtes et méthodes HTTP strictement définis dans la configuration CORS.')

# 4.11
doc.add_heading('4.11 Validation des entrées (OWASP A03)', level=2)
add_bullet('Toutes les entrées API sont validées via des modèles Pydantic v2 avec des types stricts.')
add_bullet('Validation d\'email via EmailStr.')
add_bullet('Validation conditionnelle : les champs étudiant (student_number, level_id, specialty_id) sont requis si le rôle est « etudiant » (model_validator).')
add_bullet('Les noms de table sont vérifiés contre un dictionnaire TABLE_MODELS — impossible d\'accéder à une table non autorisée.')
add_bullet('Les champs de filtre sont convertis en opérations ORM typées (eq, neq, in, gt, lt, etc.) — pas de concaténation SQL.')

# 4.12
doc.add_heading('4.12 Mécanisme de désactivation de compte', level=2)
add_bullet('L\'administrateur peut désactiver un compte utilisateur, rendant la connexion impossible.')
add_bullet('La désactivation est réversible : le compte peut être réactivé sans perte de données.')
add_bullet('Un administrateur ne peut pas se désactiver lui-même.')
add_bullet('L\'utilisateur désactivé reçoit un message explicite lors de la tentative de connexion (HTTP 403).')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. TECHNOLOGIES UTILISÉES
# ══════════════════════════════════════════════════════════════
doc.add_heading('5. Technologies Utilisées', level=1)

tech_table = doc.add_table(rows=1, cols=3)
tech_table.style = 'Light Grid Accent 1'
tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Catégorie', 'Technologie', 'Rôle']):
    tech_table.rows[0].cells[i].text = h
    for p in tech_table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

techs = [
    ('Frontend', 'React 18 + TypeScript', 'Interface utilisateur reactive'),
    ('Frontend', 'Vite', 'Bundler rapide avec HMR'),
    ('Frontend', 'TailwindCSS + shadcn/ui', 'Styling et composants UI'),
    ('Frontend', 'React Router v6', 'Navigation SPA avec protection de routes'),
    ('Frontend', 'TanStack React Query', 'Gestion du cache serveur'),
    ('Frontend', 'xlsx, jsPDF', 'Génération d\'exports Excel et PDF'),
    ('Backend', 'FastAPI (Python)', 'API REST haute performance'),
    ('Backend', 'SQLAlchemy ORM', 'Accès base de données sécurisé'),
    ('Backend', 'PostgreSQL', 'Base de données relationnelle'),
    ('Backend', 'Pydantic v2', 'Validation des données'),
    ('Backend', 'python-jose', 'Gestion des tokens JWT'),
    ('Backend', 'passlib + bcrypt', 'Hachage des mots de passe'),
    ('IA', 'Ollama', 'Serveur d\'inférence LLM local'),
    ('IA', 'Qwen2.5:3b', 'Modèle de langage pour la correction'),
    ('IA', 'FastAPI', 'Microservice de correction IA'),
]
for cat, tech, role in techs:
    row = tech_table.add_row()
    row.cells[0].text = cat
    row.cells[1].text = tech
    row.cells[2].text = role

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# 6. CONCLUSION
# ══════════════════════════════════════════════════════════════
doc.add_heading('6. Conclusion', level=1)
add_para(
    'EvalPro représente une solution complète et sécurisée pour la gestion des examens '
    'en milieu universitaire. La plateforme combine une interface utilisateur moderne et '
    'intuitive avec un backend robuste et un service d\'intelligence artificielle innovant '
    'pour la correction automatique.'
)
add_para(
    'Les multiples couches de sécurité implémentées — authentification JWT, contrôle d\'accès '
    'RBAC, protection contre les injections SQL et XSS, limitation de débit, journalisation '
    'd\'audit et isolation stricte des données — garantissent la conformité avec les meilleures '
    'pratiques de sécurité (OWASP Top 10).'
)
add_para(
    'La richesse fonctionnelle de la plateforme (gestion multi-rôles, correction IA, '
    'analytiques avancées, internationalisation, exports multiformats, notifications en temps réel) '
    'en fait un outil adapté aux besoins réels des établissements d\'enseignement supérieur, '
    'offrant un gain de temps significatif aux professeurs tout en garantissant une expérience '
    'fluide et sécurisée aux étudiants.'
)

# Add a horizontal line separator
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— Fin du document —')
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
run.italic = True

# ── Save ──
output_path = r'd:\projet api sgstock\projet_iut\EvalPro_Avantages_et_Securite.docx'
doc.save(output_path)
print(f'Document Word généré avec succès : {output_path}')
