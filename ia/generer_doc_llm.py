"""Script pour générer le document Word sur le LLM (Ollama)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ===== STYLES =====
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ===== PAGE DE TITRE =====
doc.add_paragraph()
doc.add_paragraph()
titre = doc.add_paragraph()
titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titre.add_run("LLM (Large Language Model)")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 102, 204)

sous_titre = doc.add_paragraph()
sous_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sous_titre.add_run("Explication complète et utilisation dans le projet\nSystème de Correction Automatique IA")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("Février 2026")
run.font.size = Pt(12)
run.font.italic = True

doc.add_page_break()

# ===== TABLE DES MATIÈRES =====
doc.add_heading("Table des matières", level=1)
toc_items = [
    "1. C'est quoi un LLM ?",
    "2. Quel LLM nous utilisons",
    "3. À quoi ça sert dans notre projet ?",
    "4. Comment ça marche dans le code ?",
    "   4.1 Étape 1 : Générer le prompt",
    "   4.2 Étape 2 : Appeler le LLM via Ollama",
    "   4.3 Étape 3 : Parser la réponse du LLM",
    "   4.4 Étape 4 : Appliquer le système de paliers",
    "5. Pourquoi Ollama et pas ChatGPT ?",
    "6. Le mode fallback (secours)",
    "7. Résumé : LLM vs Sentence Transformers",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ===== SECTION 1 =====
doc.add_heading("1. C'est quoi un LLM ?", level=1)

doc.add_paragraph(
    "Un LLM (Large Language Model = Grand Modèle de Langage) est une IA capable de "
    "comprendre et générer du texte comme un humain. C'est le même type de technologie "
    "que ChatGPT."
)

doc.add_paragraph(
    "La différence fondamentale avec Sentence Transformers :"
)

# Tableau comparatif fondamental
table = doc.add_table(rows=4, cols=3, style='Light List Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER

h = table.rows[0].cells
h[0].text = ""
h[1].text = "Sentence Transformers"
h[2].text = "LLM"

data = [
    ("Ce qu'il fait", "Compare deux textes (similaire ou pas ?)", "Lit, comprend, raisonne et rédige"),
    ("Analogie", "Un détecteur de plagiat", "Un correcteur humain"),
    ("Sortie", "Un nombre (0.0 à 1.0)", "Du texte libre (feedback, note, analyse)"),
]
for i, (crit, st, llm) in enumerate(data, 1):
    row = table.rows[i].cells
    row[0].text = crit
    row[1].text = st
    row[2].text = llm

doc.add_paragraph()

# ===== SECTION 2 =====
doc.add_heading("2. Quel LLM nous utilisons", level=1)

doc.add_paragraph("Configuration définie dans le fichier ia_correction/config.py :")

p = doc.add_paragraph()
run = p.add_run('OLLAMA_MODEL = "llama3.2"\nOLLAMA_HOST = "http://localhost:11434"')
run.font.name = 'Consolas'
run.font.size = Pt(10)
run.bold = True

doc.add_paragraph()

items = [
    ("Llama 3.2", "Modèle open source développé par Meta (Facebook). C'est un modèle performant capable de comprendre et générer du texte en français."),
    ("Ollama", "Logiciel qui fait tourner le LLM localement sur votre PC. Pas besoin d'internet, pas de coût d'API. Les données restent sur la machine."),
]
for mot, explication in items:
    p = doc.add_paragraph()
    run = p.add_run(f"{mot} : ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(explication)

# ===== SECTION 3 =====
doc.add_heading("3. À quoi ça sert dans notre projet ?", level=1)

doc.add_paragraph(
    "Le LLM est utilisé dans le fichier ia_correction/correcteurs/longue_corrector.py "
    "pour corriger les questions à réponse longue (dissertations, explications, argumentations)."
)

doc.add_heading("Pourquoi un LLM est nécessaire ici ?", level=2)

doc.add_paragraph(
    "Sentence Transformers ne suffit pas pour les réponses longues car :"
)

raisons = [
    "Il faut évaluer la qualité de l'argumentation",
    "Il faut vérifier plusieurs critères (pertinence, structure, connaissances)",
    "Il faut générer un feedback personnalisé pour l'élève",
    "La réponse peut être correcte même avec des formulations très différentes",
]
for r in raisons:
    doc.add_paragraph(r, style='List Bullet')

doc.add_paragraph()
doc.add_heading("Exemple concret", level=2)

p = doc.add_paragraph()
run = p.add_run("Question : ")
run.bold = True
p.add_run('"Expliquez le modèle TCP/IP et ses 4 couches"')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Réponse de l'élève : ")
run.bold = True
p.add_run(
    '"Le modèle TCP/IP comporte 4 couches : la couche application (HTTP, FTP), '
    'la couche transport (TCP, UDP), la couche internet (IP) et la couche accès réseau."'
)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Ce que le LLM doit faire :")
run.bold = True

evaluations = [
    "1. Vérifier que les 4 couches sont mentionnées → ✅",
    "2. Vérifier les protocoles cités → ✅",
    "3. Évaluer la clarté de l'explication → ✅",
    '4. Donner une note : 4.5/5',
    '5. Rédiger un feedback : "Bonne réponse, il manque un exemple de protocole couche accès réseau"',
]
for ev in evaluations:
    doc.add_paragraph(ev, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Important : ")
run.bold = True
run.font.color.rgb = RGBColor(204, 0, 0)
p.add_run(
    "Sentence Transformers ne pourrait jamais faire ça — il donnerait juste un score "
    "de similarité sans aucun détail."
)

# ===== SECTION 4 =====
doc.add_heading("4. Comment ça marche dans le code ?", level=1)

doc.add_paragraph(
    "Le processus se déroule en 4 étapes dans le fichier longue_corrector.py (lignes 96 à 170) :"
)

# ÉTAPE 1
doc.add_heading("4.1 Étape 1 : Générer le prompt (l'instruction au LLM)", level=2)

doc.add_paragraph(
    "Le prompt est le texte qu'on envoie au LLM pour lui dire quoi faire. "
    "C'est comme donner des consignes à un correcteur humain :"
)

code1 = '''prompt = f"""Note: 0-{points_max}
Q: {question_courte}
R: {reponse_courte}

Format:
NOTE: X/{points_max}
FEEDBACK: (1 ligne)
RÉPONSE: (essentiel)"""'''

p = doc.add_paragraph()
run = p.add_run(code1)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(30, 30, 30)

doc.add_paragraph()
doc.add_paragraph(
    'Le prompt dit au LLM : "Tu es un correcteur, voici la question, voici la réponse '
    "de l'élève, donne une note et un feedback.\""
)

# ÉTAPE 2
doc.add_heading("4.2 Étape 2 : Appeler le LLM via Ollama", level=2)

code2 = '''response = ollama.chat(
    model='llama3.2',
    messages=[
        {'role': 'system', 'content': 'Correcteur. Réponds brièvement.'},
        {'role': 'user', 'content': prompt}
    ],
    options={
        'temperature': 0.1,   # Très déterministe (pas de créativité)
        'num_predict': 150,   # Maximum 150 tokens en sortie
        'top_k': 10,          # Limite les choix de mots
    }
)'''

p = doc.add_paragraph()
run = p.add_run(code2)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(30, 30, 30)

doc.add_paragraph()
doc.add_heading("Explication des paramètres importants :", level=3)

params = [
    ("temperature: 0.1", "Le LLM donne toujours la même réponse (pas de hasard). Comme un correcteur rigoureux. Une température élevée (ex: 0.9) rendrait les réponses créatives et variables — pas souhaitable pour une correction."),
    ("num_predict: 150", "Limite la longueur de la réponse à 150 tokens maximum pour accélérer le traitement. Un feedback de correction n'a pas besoin d'être long."),
    ("top_k: 10", "Réduit le nombre de mots candidats à chaque étape de génération, ce qui accélère le calcul tout en maintenant la qualité."),
]
for param, expl in params:
    p = doc.add_paragraph()
    run = p.add_run(f"{param} → ")
    run.bold = True
    run.font.color.rgb = RGBColor(0, 102, 204)
    p.add_run(expl)

# ÉTAPE 3
doc.add_heading("4.3 Étape 3 : Parser (analyser) la réponse du LLM", level=2)

doc.add_paragraph(
    "Le LLM répond en texte libre. Il faut extraire la note et le feedback "
    "avec des expressions régulières (regex) :"
)

code3 = '''# Le LLM répond : "NOTE: 3.5/5\\nFEEDBACK: Bonne réponse..."

# Extraire la note avec une regex
match = re.search(r'NOTE:\\s*(\\d+\\.?\\d*)\\s*/\\s*(\\d+)', reponse_llm)
points = float(match.group(1))  # → 3.5'''

p = doc.add_paragraph()
run = p.add_run(code3)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(30, 30, 30)

doc.add_paragraph()
doc.add_paragraph(
    "Le parser gère aussi les cas où le LLM ne respecte pas le format demandé. "
    "Il utilise des mots-clés comme \"excellent\", \"bien\", \"moyen\" pour estimer la note "
    "si le format structuré n'est pas détecté."
)

# ÉTAPE 4
doc.add_heading("4.4 Étape 4 : Appliquer le système de paliers", level=2)

doc.add_paragraph(
    "Au lieu de garder la note brute du LLM, on applique un système de paliers "
    "pour une notation standardisée :"
)

table_paliers = doc.add_table(rows=5, cols=3, style='Light List Accent 1')
table_paliers.alignment = WD_TABLE_ALIGNMENT.CENTER

h = table_paliers.rows[0].cells
h[0].text = "Score LLM brut"
h[1].text = "Points attribués"
h[2].text = "Appréciation"

paliers = [
    ("> 50%", "100% des points", "Réponse correcte"),
    ("30% - 50%", "50% des points", "Partiellement correcte"),
    ("25% - 30%", "25% des points", "Réponse incomplète"),
    ("< 25%", "0 point", "Réponse insuffisante"),
]
for i, (score, pts, app) in enumerate(paliers, 1):
    row = table_paliers.rows[i].cells
    row[0].text = score
    row[1].text = pts
    row[2].text = app

doc.add_paragraph()

code4 = '''pourcentage_reponse = (resultat_analyse['points_obtenus'] / question.points_max) * 100

if pourcentage_reponse > 50:
    points_finaux = question.points_max          # Point complet
    appreciation_palier = "Réponse correcte"
elif pourcentage_reponse >= 30:
    points_finaux = question.points_max * 0.5    # Moitié
    appreciation_palier = "Réponse partiellement correcte"
elif pourcentage_reponse >= 25:
    points_finaux = question.points_max * 0.25   # Quart
    appreciation_palier = "Réponse incomplète"
else:
    points_finaux = 0                             # Zéro
    appreciation_palier = "Réponse insuffisante"'''

p = doc.add_paragraph()
run = p.add_run(code4)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(30, 30, 30)

# ===== SECTION 5 =====
doc.add_heading("5. Pourquoi Ollama et pas ChatGPT ?", level=1)

doc.add_paragraph(
    "On aurait pu utiliser l'API de ChatGPT (OpenAI) pour la correction. "
    "Voici pourquoi nous avons choisi Ollama avec un modèle local :"
)

table_vs = doc.add_table(rows=7, cols=3, style='Light List Accent 1')
table_vs.alignment = WD_TABLE_ALIGNMENT.CENTER

h = table_vs.rows[0].cells
h[0].text = "Critère"
h[1].text = "Ollama (local)"
h[2].text = "ChatGPT (API)"

comparaisons = [
    ("Coût", "Gratuit", "~0.01€ par question"),
    ("Internet", "Pas nécessaire", "Obligatoire"),
    ("Vie privée", "Données restent sur le PC", "Données envoyées à OpenAI"),
    ("Vitesse", "~2-5s (dépend du PC)", "~1-2s"),
    ("Disponibilité", "Toujours (même hors-ligne)", "Dépend des serveurs"),
    ("Copies d'examen", "Confidentialité garantie", "Risque de fuite"),
]
for i, (crit, oll, gpt) in enumerate(comparaisons, 1):
    row = table_vs.rows[i].cells
    row[0].text = crit
    row[1].text = oll
    row[2].text = gpt

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Le choix d'Ollama est stratégique : ")
run.bold = True
run.font.color.rgb = RGBColor(0, 102, 204)
p.add_run(
    "les copies d'élèves sont des données sensibles qu'on ne veut pas envoyer sur Internet. "
    "Avec Ollama, tout reste en local, la confidentialité est garantie et il n'y a aucun coût."
)

# ===== SECTION 6 =====
doc.add_heading("6. Le mode fallback (secours)", level=1)

doc.add_paragraph(
    "Si Ollama n'est pas installé ou n'est pas démarré, le système ne plante pas. "
    "Il bascule automatiquement sur une correction basique (lignes 65-86 du fichier longue_corrector.py) :"
)

code_fallback = '''if not self.ollama_disponible:
    return self._corriger_basique(question, reponse)
    # → Score basé uniquement sur la longueur de la réponse
    # → Maximum 50% des points (car pas fiable)'''

p = doc.add_paragraph()
run = p.add_run(code_fallback)
run.font.name = 'Consolas'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(30, 30, 30)

doc.add_paragraph()
doc.add_paragraph(
    "La correction basique ne donne jamais plus de 50% des points car elle est peu fiable. "
    "Elle est là uniquement pour que le système fonctionne même sans LLM, "
    "en attendant que l'enseignant puisse corriger manuellement."
)

# ===== SECTION 7 =====
doc.add_heading("7. Résumé : LLM vs Sentence Transformers dans le projet", level=1)

doc.add_paragraph(
    "Voici le schéma récapitulatif du choix de technologie selon le type de question :"
)

schema = """
┌──────────────────────────────────────────────────────────┐
│           SYSTÈME DE CORRECTION IA                        │
├──────────────────────────────────────────────────────────┤
│                                                          │
│  Questions COURTES ──→ Sentence Transformers             │
│  "Quel protocole       (comparaison sémantique)          │
│   utilise le port 80?" → Score: 0.92 → ✅ Correct       │
│                         Temps: ~100ms                    │
│                                                          │
│  Questions LONGUES ──→ LLM (Ollama / Llama3.2)          │
│  "Expliquez le         (lecture + raisonnement           │
│   modèle TCP/IP"       + rédaction de feedback)          │
│                        → Note: 4/5                       │
│                        → Feedback personnalisé           │
│                         Temps: ~2-5s                     │
│                                                          │
└──────────────────────────────────────────────────────────┘
"""

p = doc.add_paragraph()
run = p.add_run(schema)
run.font.name = 'Consolas'
run.font.size = Pt(9)

doc.add_paragraph()

doc.add_heading("Phrase de synthèse pour l'exposé", level=2)

p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
p.paragraph_format.right_indent = Cm(1)
run = p.add_run(
    "« Pour les questions longues, on utilise un LLM (Llama 3.2) via Ollama qui fonctionne "
    "localement. Contrairement à Sentence Transformers qui compare juste deux textes, le LLM "
    "lit la réponse de l'élève, évalue sa qualité sur plusieurs critères et rédige un feedback "
    "personnalisé, exactement comme le ferait un correcteur humain — mais en 3 secondes au lieu "
    "de 5 minutes. »"
)
run.italic = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0, 102, 204)

# ===== SAUVEGARDER =====
output_path = r"d:\projet api sgstock\projet_iut\ia\LLM_Ollama_Explication.docx"
doc.save(output_path)
print(f"✅ Document Word créé avec succès : {output_path}")
